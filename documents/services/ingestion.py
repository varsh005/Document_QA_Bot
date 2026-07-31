"""
Ingestion service.

Responsible for turning an uploaded file (PDF or DOCX) into plain text,
then splitting that text into overlapping chunks ready for embedding.

Design note: both file formats converge into the same plain-text output.
Everything downstream (chunking, embeddings, retrieval) is format-agnostic.
"""
import pdfplumber
import docx
import tiktoken
from django.conf import settings


class UnsupportedFileTypeError(Exception):
    pass


class EmptyDocumentError(Exception):
    pass


def extract_text(file_path: str, file_type: str) -> str:
    """Extract raw text from a PDF or DOCX file on disk."""
    if file_type == 'pdf':
        return _extract_pdf_text(file_path)
    elif file_type == 'docx':
        return _extract_docx_text(file_path)
    raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")


def _extract_pdf_text(file_path: str) -> str:
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            if text.strip():
                pages.append(text)
    return '\n\n'.join(pages)


def _extract_docx_text(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables often hold important structured notes (e.g. comparison tables);
    # pull their text in too so it isn't silently dropped.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(' | '.join(cells))

    return '\n\n'.join(paragraphs)


def chunk_text(text: str, chunk_size_tokens: int = None, overlap_tokens: int = None) -> list[str]:
    """
    Split text into overlapping chunks measured in tokens (not characters),
    since that's what the embedding and chat models actually bill and limit on.

    Overlap exists so a sentence that gets cut at a chunk boundary still
    appears in full in the neighboring chunk - otherwise retrieval can miss
    an answer that straddles two chunks.
    """
    chunk_size_tokens = chunk_size_tokens or settings.CHUNK_SIZE_TOKENS
    overlap_tokens = overlap_tokens or settings.CHUNK_OVERLAP_TOKENS

    text = text.strip()
    if not text:
        raise EmptyDocumentError(
            "No extractable text was found in this file. "
            "It may be a scanned/image-only document, which isn't supported yet."
        )

    encoder = tiktoken.get_encoding('cl100k_base')
    tokens = encoder.encode(text)

    chunks = []
    start = 0
    step = chunk_size_tokens - overlap_tokens
    while start < len(tokens):
        end = min(start + chunk_size_tokens, len(tokens))
        chunk_tokens = tokens[start:end]
        chunks.append(encoder.decode(chunk_tokens))
        if end == len(tokens):
            break
        start += step

    return chunks


def process_upload(file_path: str, file_type: str) -> list[str]:
    """Full ingestion pipeline: extract text, then chunk it. Returns list of chunk strings."""
    text = extract_text(file_path, file_type)
    return chunk_text(text)
